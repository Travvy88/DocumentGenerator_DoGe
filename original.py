import colorsys
import io
import json
import os
import random
import re
import string
import sys
import time
from io import BytesIO
from pathlib import Path
from urllib.parse import urlparse, urljoin

import cv2
import matplotlib
import matplotlib.font_manager
import numpy as np
import requests
import torch
from PIL import ImageDraw
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_COLOR_INDEX, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import CT_RPr, parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.oxml.xmlchemy import XmlString
from docx.shared import Pt, RGBColor
from pdf2image import convert_from_path, convert_from_bytes
from tqdm import tqdm
from unoserver import client


class WikiDatasetGenerator:
    def __init__(self,
                 output_folder,
                 port=2003,
                 dataset_config=None,
                 resize=224,
                 num_colors=20000,
                 debug_mode=False,
                 seed=88
                 ):

        if dataset_config is None:
            dataset_config = dict(
                p_2columns=0.0,
                font_size=(5, 15),
                heading_relative_size=(1, 3),
                heading_bold_p=0.5,
                heading_center_alignment_p=0.5,
                table_max_rows=20,
                table_max_cols=6
            )

        # main data and properties
        #self.srv = server.UnoServer(port=port, uno_port="2202")
        #self.srv.start()
        #os.system("")
        self.conv = client.UnoClient(port=port)

        self.dataset_config = dataset_config
        self.colors = self.__init_different_colors(num_colors)
        self.image_counter = 0
        self.output_folder = output_folder
        self.fonts_names = self.__list_available_fonts()
        self.debug_mode = debug_mode
        self.resize = resize
        self.urls = []

        # current doc properties
        self.font_size = None
        self.font_name = None
        self.num_columns = None
        self.heading_bold = None
        self.heading_center_alignment = None
        self.heading_relative_size = None

        # current doc data
        self.doc = None
        self.images = []
        self.color_ptr = 0
        self.color2word = {}
        self.annotations = []

        random.seed(seed)
        os.environ['PYTHONHASHSEED'] = str(seed)
        np.random.seed(seed)
        torch.manual_seed(seed)
        torch.cuda.manual_seed(seed)
        torch.backends.cudnn.deterministic = True

    def __list_available_fonts(self):
        font_paths = matplotlib.font_manager.findSystemFonts(fontpaths=None, fontext='ttf')
        font_names = set()

        for font_path in font_paths:
            try:
                font = matplotlib.font_manager.get_font(font_path)
                font_names.add(font.family_name)
            except RuntimeError as e:
                print(f"Could not load font from path: {font_path}, error: {e}")

        return list(font_names)

    def init_doc(self):
        self.doc = Document()
        self.color_ptr = 0
        self.num_columns = 1
        if np.random.binomial(1, self.dataset_config["p_2columns"]):
            self.num_columns = 2

        self.font_size = np.random.randint(*self.dataset_config["font_size"])
        self.font_name = np.random.choice(self.fonts_names)

        self.heading_bold = bool(np.random.binomial(1, self.dataset_config["heading_bold_p"]))
        self.heading_relative_size = np.random.uniform(*self.dataset_config["heading_relative_size"])
        self.heading_center_alignment = np.random.binomial(1, self.dataset_config["heading_center_alignment_p"])

    def __init_different_colors(self, num_colors):
        colors = []
        hex_color = "#000000"
        x = int(num_colors ** (1/3)) + 1
        for i in range(x):
            for j in range(x):
                for k in range(x):
                    # Convert HEX to RGB
                    r = int(hex_color[1:3], 16)
                    g = int(hex_color[3:5], 16)
                    b = int(hex_color[5:7], 16)

                    # Increment RGB values
                    r = (r + i) % 256
                    g = (g + j) % 256
                    b = (b + k) % 256
                    hex_color = '#{:02x}{:02x}{:02x}'.format(r, g, b)
                    colors.append(hex_color)
        return colors

    def get_color_and_save_word_to_annotation(self, word, run):
        try:
            color = self.colors[self.color_ptr]
        except:
            raise Exception(f"Need more colors, {len(self.colors)} not enough")

        self.color_ptr += 1
        self.color2word[color] = word
        tag = run._r

        # Create XML element
        shd = OxmlElement('w:shd')

        # Add attributes to the element
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color)

        # Set the font size - this is important! Without this step the
        # tag.rPr value below will be None.
        run.element.get_or_add_rPr()

        tag.rPr.append(shd)

        run.font.color.rgb = RGBColor(*tuple(int(color[i:i + 2], 16) for i in (1, 3, 5)))

    def convert_to_uncolored_docx(self):
        for paragraph in self.doc.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
                rpr = run.element.get_or_add_rPr()
                element = rpr.find(qn('w:shd'))
                if element is not None:
                    element.set(qn('w:fill'), "#FFFFFF")

        for table in self.doc.tables:
            self.set_table_border_white(table, "000000")
            for row in table.rows:
                for cell in row.cells:
                    for run in cell.paragraphs[0].runs:
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        rpr = run.element.get_or_add_rPr()
                        element = rpr.find(qn('w:shd'))
                        if element is not None:
                            element.set(qn('w:fill'), "#FFFFFF")

    def add_text(self, text, paragraph, formatting=None, prev_word=" "):
        text = re.sub(r'\[.*?\]', '', text)
        for word in re.split(r'\s+', text):
            if word:
                if word[0] not in ",.?!:;)»" and prev_word[0] not in "«":
                    paragraph.add_run(' ')

                run = paragraph.add_run(word)
                self.get_color_and_save_word_to_annotation(word, run)
                if formatting == 'b':
                    run.bold = True
                if formatting == 'i':
                    run.italic = True
                if formatting == 'u':
                    run.underline = True
                prev_word = word
                run.font.size = Pt(self.font_size)
                run.font.name = self.font_name
        return prev_word

    def add_heading(self, text, heading_level):
        if text not in ["Contents"]:
            paragraph = self.doc.add_heading(level=heading_level)
            self.add_text(text, paragraph)
            for run in paragraph.runs:
                run.font.size = Pt(self.heading_relative_size * self.font_size)
            if self.heading_center_alignment:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def add_table(self, element):
        rows = element.find_all('tr')
        parsed_table = []
        for row in rows:
            cells = row.find_all(['th', 'td'])
            parsed_row = []
            for cell in cells:
                parsed_row.append(cell.text.strip())
            parsed_table.append(parsed_row)

        rows = len(parsed_table)
        cols = max(len(row) for row in parsed_table)
        if rows <= self.dataset_config["table_max_rows"] and cols <= self.dataset_config["table_max_cols"]:
            table = self.doc.add_table(rows=len(parsed_table), cols=max(len(row) for row in parsed_table))
            table.style = 'TableGrid'
            self.set_table_border_white(table, "FFFFFF")
            # Populating table data
            for i, row_data in enumerate(parsed_table):
                for j, cell_data in enumerate(row_data):
                    self.add_text(cell_data, table.cell(i, j).paragraphs[0])
                    '''table.cell(i, j).paragraphs[0].paragraph_format.line_spacing = Pt(24)
                    for run in table.cell(i, j).paragraphs[0].runs:
                        run.font.size = self.doc_config["font_size"]'''

    def set_table_border_white(self, table, color):
        tbl = table._element
        tbl_pr = tbl.tblPr

        # Create a new border element
        tbl_borders = OxmlElement('w:tblBorders')

        # Create a list of border attributes
        borders = [
            'top',
            'left',
            'bottom',
            'right',
            'insideH',
            'insideV'
        ]

        # Iterate through each border attribute
        for border in borders:
            border_element = OxmlElement(f'w:{border}')
            border_element.set(qn('w:val'), 'single')
            border_element.set(qn('w:sz'), '4')
            border_element.set(qn('w:space'), '0')
            border_element.set(qn('w:color'), color)
            tbl_borders.append(border_element)

        tbl_pr.append(tbl_borders)

    def configure_several_columns(self, num_columns):
        section = self.doc.sections[0]
        sectPr = section._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set(qn('w:num'), str(num_columns))

    def create_colored_docx(self, url):
        self.doc.add_paragraph()
        response = requests.get(url)
        if response.status_code == 200:
            html_content = response.text
            soup = BeautifulSoup(html_content, 'html.parser')

            if self.num_columns != 1:
                self.configure_several_columns(self.num_columns)

            # Process each element in the page
            for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', "table"]):
                paragraph = self.doc.add_paragraph()
                paragraph.paragraph_format.space_after = 0
                if element.name.startswith('h') and "Heading" not in self.doc.paragraphs[-2].style.style_id:
                    self.add_heading(element.text, int(element.name[1]))
                elif element.name == "table":
                    self.add_table(element)
                else:
                    prev_word = " "
                    prev_word = self.add_text("\t    ", paragraph, prev_word=prev_word)  # TODO
                    for child in element.children:
                        # TODO нет форматирования
                        prev_word = self.add_text(child.get_text(), paragraph, formatting=child.name, prev_word=prev_word)
        else:
            print("Error: Unable to fetch Wikipedia page.")
            return None

    def docx2imgs(self, temp_files_name='doc'):
        #self.doc.save(f'{self.output_folder}/{temp_files_name}.docx')

        out = io.BytesIO()
        self.doc.save(out)
        doc_bytes = out.getvalue()
        pdf_bytes = self.conv.convert(indata=doc_bytes, convert_to='pdf')
        self.images = convert_from_bytes(pdf_bytes)

        #os.system(f"soffice --headless --convert-to pdf {self.output_folder}/{temp_files_name}.docx --outdir {self.output_folder}")
        #self.images = convert_from_path(f"{self.output_folder}/{temp_files_name}.pdf")

        #if not self.debug_mode:
            #os.remove(f"{self.output_folder}/{temp_files_name}.docx")
            #os.remove(f"{self.output_folder}/{temp_files_name}.pdf")

    def get_bboxes(self):
        self.annotations = []
        for count, image_pil in enumerate(self.images):
            draw = ImageDraw.Draw(image_pil)
            width, height = image_pil.size
            image_annotations = []
            image = np.asarray(image_pil)

            thr = cv2.cvtColor(image, cv2.COLOR_RGB2GRAY)
            thr = cv2.threshold(thr, 254, 255, cv2.THRESH_BINARY_INV)[1]
            cnts = cv2.findContours(thr, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)[0]

            for c in cnts:
                peri = cv2.arcLength(c, True)
                approx = cv2.approxPolyDP(c, 0.015 * peri, True)

                if len(approx) == 4:
                    x, y, w, h = cv2.boundingRect(approx)
                    if self.debug_mode:
                        bbox = (x, y, x + w, y + h)
                        draw.rectangle(bbox, outline="red")

                    rgb_color = image_pil.getpixel((x+1, y+1))
                    color = '#%02x%02x%02x' % (rgb_color)
                    if color in self.color2word:
                        word = self.color2word[color]
                        image_annotations.append(((x / width, y / height, w / width, h / height), word))
            self.annotations.append(image_annotations)
            if self.debug_mode:
                image_pil.save(f"countours.png")

    def save_images_and_annotations(self):
        for i, image in enumerate(self.images):
            image.resize((self.resize, self.resize)).save(Path(self.output_folder) / f"im_{self.image_counter}.png")
            with open(Path(self.output_folder) / f"im_{self.image_counter}.png.json", "w") as f:
                json.dump(self.annotations[i], f)
            self.image_counter += 1

    def pipeline(self, url):
        self.init_doc()
        self.create_colored_docx(url)
        self.docx2imgs(temp_files_name="colored")
        if self.debug_mode:
            for i, image in enumerate(self.images):
                image.save(f"im_colored_{i}.png")
        self.get_bboxes()
        self.convert_to_uncolored_docx()
        self.docx2imgs(temp_files_name="final")
        self.save_images_and_annotations()

    def generate(self):
        print("Generation dataset...")
        for url in tqdm(self.urls):
            try:
                self.pipeline(url)
            except Exception as e:
                print(e)

    def is_valid_url(self, url):
        # Check if the URL is a valid Wikipedia article URL
        parsed = urlparse(url)
        if parsed.scheme in ('http', 'https') and 'wikipedia.org' in parsed.netloc:
            path = parsed.path
            if path.startswith('/wiki/') and not any(sub in path for sub in [':', '/wiki/Main_Page']):
                return True
        return False

    def parse_urls(self, start_url, max_urls):
        ptr = 0
        self.urls.append(start_url)

        pbar = tqdm(initial=1, total=max_urls)
        while len(self.urls) < max_urls:
            url = self.urls[ptr]
            try:
                response = requests.get(url)
                response.raise_for_status()
            except requests.exceptions.RequestException as e:
                print(f"Failed to retrieve {url}: {e}")
                return

            # Parse the page content
            soup = BeautifulSoup(response.content, 'html.parser')

            # Find all links on the page
            links = soup.find_all('a', href=True)
            for link in links:
                href = link['href']
                full_url = urljoin(url, href)
                if self.is_valid_url(full_url) and full_url not in self.urls and len(self.urls) < max_urls:
                    self.urls.append(full_url)
                    pbar.update(1)
            ptr += 1
        return self.urls

