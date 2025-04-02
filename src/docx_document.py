import cProfile
import io
import re
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
import matplotlib.font_manager
import numba
import numpy as np
from pdf2image import convert_from_bytes
from PIL import Image
def profileit(func):
    def wrapper(*args, **kwargs):
        datafn = func.__name__ + ".profile" # Name the data file sensibly
        prof = cProfile.Profile()
        retval = prof.runcall(func, *args, **kwargs)
        prof.dump_stats(datafn)
        return retval

    return wrapper

class DocxDocument:
    def __init__(self, docx_config, uno_client):
        self.docx_config = docx_config
        self.uno_client = uno_client

        self.doc = Document()
        self.colors = self._init_colors(docx_config["max_words"])

        self.color2word = {}
        self.color_ptr = 0

        # sample random settings from docx_config 
        if np.random.binomial(1, self.docx_config["p_2columns"]):
            self.num_columns = 2
        else:
            self.num_columns = 1
        self.configure_several_columns()

        self.font_size = Pt(np.random.randint(*self.docx_config["font_size_interval"]))
        self.font_name = np.random.choice(self._list_available_fonts())
        
        self.line_spacing = np.random.choice(
            (WD_LINE_SPACING.ONE_POINT_FIVE, WD_LINE_SPACING.DOUBLE), 
            p=self._normalize_probabilities(self.docx_config["p_line_spacing"])) 

        self.paragraph_alignment = np.random.choice(
            (WD_PARAGRAPH_ALIGNMENT.CENTER, WD_PARAGRAPH_ALIGNMENT.LEFT, 
             WD_PARAGRAPH_ALIGNMENT.RIGHT, WD_PARAGRAPH_ALIGNMENT.JUSTIFY), 
            p=self._normalize_probabilities(self.docx_config["p_text_alignment"]))
    
        self.heading_bold = bool(np.random.binomial(1, self.docx_config["p_heading_bold"]))
        self.heading_relative_size = np.random.uniform(*self.docx_config["heading_relative_size_interval"])
        self.heading_size = Pt(self.heading_relative_size * self.font_size)
        self.heading_alignment = np.random.choice(
            (WD_PARAGRAPH_ALIGNMENT.CENTER, WD_PARAGRAPH_ALIGNMENT.LEFT, 
             WD_PARAGRAPH_ALIGNMENT.RIGHT, WD_PARAGRAPH_ALIGNMENT.JUSTIFY), 
            p=self._normalize_probabilities(self.docx_config["p_heading_alignment"]))
    
    def _normalize_probabilities(self, p):
        return np.array(p) / sum(p)

    def _list_available_fonts(self):
        font_paths = matplotlib.font_manager.findSystemFonts(fontpaths=None, fontext='ttf')
        font_names = set()

        for font_path in font_paths:
            try:
                font = matplotlib.font_manager.get_font(font_path)
                font_names.add(font.family_name)
            except RuntimeError as e:
                print(f"Could not load font from path: {font_path}, error: {e}")
        return list(font_names)
    
    def _init_colors(self, max_colors):
        colors = []
        hex_color = "#000000"
        x = int(max_colors ** (1/3)) + 1
        for i in range(x):
            for j in range(x):
                for k in range(x):
                    # Convert HEX to RGB
                    r = int(hex_color[1:3], 16)
                    g = int(hex_color[3:5], 16)
                    b = int(hex_color[5:7], 16)

                    # Increment RGB values
                    r = (r + i) % 256
                    g = (g + j) % 256
                    b = (b + k) % 256
                    hex_color = '#{:02x}{:02x}{:02x}'.format(r, g, b)
                    colors.append(hex_color)
        return colors
    
    def configure_several_columns(self):
        section = self.doc.sections[0]
        sectPr = section._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set(qn('w:num'), str(self.num_columns))    

    def add_paragraph(self):
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = self.paragraph_alignment
        paragraph.paragraph_format.space_after = 0
        paragraph.paragraph_format.line_spacing_rule = self.line_spacing
        
        return paragraph

    def add_heading(self, element):
        text = element.text
        level = int(element.name[1])
        
        if len(self.doc.paragraphs) > 1 and "Heading" in self.doc.paragraphs[-2].style.style_id:
            return
        
        if text not in ["Contents"]:
            paragraph = self.doc.add_heading(level=level)
            _, metadata = self.add_words(text, paragraph)
            self.color2word.update(metadata)
            #for run in paragraph.runs:
                #run.font.size = self.heading_size
            paragraph.alignment = self.heading_alignment
            
            p = self.add_paragraph()
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    
    def add_table(self, html_element):
        rows = html_element.find_all('tr')
        parsed_table = []
        for row in rows:
            cells = row.find_all(['th', 'td'])
            parsed_row = []
            for cell in cells:
                parsed_row.append(cell.text.strip())
            parsed_table.append(parsed_row)

        rows = len(parsed_table)
        cols = max(len(row) for row in parsed_table)
        if rows <= self.docx_config["table_max_rows"] and cols <= self.docx_config["table_max_cols"]:
            table = self.doc.add_table(rows=len(parsed_table), cols=max(len(row) for row in parsed_table))
            table.style = 'TableGrid'
            self.set_table_border_color(table, "FFFFFF")
            # Populating table data
            for i, row_data in enumerate(parsed_table):
                for j, cell_data in enumerate(row_data):
                    _, metadata = self.add_words(cell_data, table.cell(i, j).paragraphs[0])
                    self.color2word.update(metadata)
                    '''table.cell(i, j).paragraphs[0].paragraph_format.line_spacing = Pt(24)
                    for run in table.cell(i, j).paragraphs[0].runs:
                        run.font.size = self.doc_config["font_size"]'''
    def add_text(self, html_element):
        paragraph = self.add_paragraph()
        prev_word = " "
        first_word = True
        for i, child in enumerate(html_element.children):
            if i > 0:
                first_word = False
            prev_word, metadata = self.add_words(child.get_text(), paragraph, formatting=child.name, prev_word=prev_word, first_word=first_word)
            self.color2word.update(metadata)

    def add_words(self, text, paragraph, formatting=None, prev_word=" ", first_word=False):
        text = re.sub(r'\[.*?\]', '', text)
        words = re.split(r'\s+', text)
        metadata = {}
        for word in words:
            if word:
                if word[0] not in ",.?!:;)}]»" and prev_word[-1] not in "«[{(":
                    if first_word:
                        paragraph.add_run(' ' * 4)
                    else:
                        paragraph.add_run(' ')

                run = paragraph.add_run(word)
                color = self.color_word(run)
                metadata[color] = word
                if formatting == 'b':
                    run.bold = True
                if formatting == 'i':
                    run.italic = True
                if formatting == 'u':
                    run.underline = True
                prev_word = word
                run.font.size = self.font_size
                run.font.name = self.font_name
        return prev_word, metadata
    
    def color_word(self, run):
        color = self.colors[self.color_ptr]

        self.color_ptr += 1
        tag = run._r

        # Create XML element
        shd = OxmlElement('w:shd')

        # Add attributes to the element
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color)

        # Set the font size - this is important! Without this step the
        # tag.rPr value below will be None.
        run.element.get_or_add_rPr()

        tag.rPr.append(shd)

        run.font.color.rgb = RGBColor(*tuple(int(color[i:i + 2], 16) for i in (1, 3, 5)))
        return color

    def set_table_border_color(self, table, color):
        tbl = table._element
        tbl_pr = tbl.tblPr

        # Create a new border element
        tbl_borders = OxmlElement('w:tblBorders')

        # Create a list of border attributes
        borders = [
            'top',
            'left',
            'bottom',
            'right',
            'insideH',
            'insideV'
        ]

        # Iterate through each border attribute
        for border in borders:
            border_element = OxmlElement(f'w:{border}')
            border_element.set(qn('w:val'), 'single')
            border_element.set(qn('w:sz'), '4')
            border_element.set(qn('w:space'), '0')
            border_element.set(qn('w:color'), color)
            tbl_borders.append(border_element)

        tbl_pr.append(tbl_borders)

    def save_docx(self, path):
        self.doc.save(path)

    #@profileit
    def get_images(self, image_size, dpi) -> list[Image]:
        out = io.BytesIO()
        self.doc.save(out)
        doc_bytes = out.getvalue()
        pdf_bytes = self.uno_client.convert(indata=doc_bytes, convert_to='pdf')
        return convert_from_bytes(pdf_bytes, dpi=dpi, size=image_size) 
    
    def convert_to_uncolored_docx(self):
        for paragraph in self.doc.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
                rpr = run.element.get_or_add_rPr()
                element = rpr.find(qn('w:shd'))
                if element is not None:
                    element.set(qn('w:fill'), "#FFFFFF")

        for table in self.doc.tables:
            self.set_table_border_color(table, "000000")
            for row in table.rows:
                for cell in row.cells:
                    for run in cell.paragraphs[0].runs:
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        rpr = run.element.get_or_add_rPr()
                        element = rpr.find(qn('w:shd'))
                        if element is not None:
                            element.set(qn('w:fill'), "#FFFFFF")
        
    def get_num_words(self):
        return len(self.color2word)
