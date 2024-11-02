import io
from pdf2image import convert_from_bytes
from unoserver import client
import re
from bs4 import BeautifulSoup
from docx import Document
import requests
from docx.shared import Pt

doc = Document()
def create_colored_docx(doc, url):
        doc.add_paragraph()
        response = requests.get(url)
        if response.status_code == 200:
            html_content = response.text
            soup = BeautifulSoup(html_content, 'html.parser')

            # Process each element in the page
            for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', "table"]):
                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.space_after = 0
                first_word = True
                prev_word = " "
                for i, child in enumerate(element.children):
                    if i > 0:
                        first_word = False
                    # TODO нет форматирования
                    prev_word = add_text(child.get_text(), paragraph, formatting=child.name, prev_word=prev_word, first_word=first_word)
        else:
            print("Error: Unable to fetch Wikipedia page.")
            return None
        
        return doc

def add_text_(text, paragraph, formatting=None, first_word=False):
    #text = re.sub(r'\[.*?\]', '', text)
    text_tokens = text.split()
    print(text, '==', text_tokens)
    for word in text_tokens:
            if first_word:
                paragraph.add_run('\t')
            else:
                paragraph.add_run(' ')

            run = paragraph.add_run(word)
            #self.get_color_and_save_word_to_annotation(word, run)
            if formatting == 'b':
                run.bold = True
            if formatting == 'i':
                run.italic = True
            if formatting == 'u':
                run.underline = True
            prev_word = word
            #run.font.size = Pt(self.font_size)
            #run.font.name = self.font_name
    

  
def add_text(text, paragraph, formatting=None, prev_word=" ", first_word=False):
    text = re.sub(r'\[.*?\]', '', text)
    for word in re.split(r'\s+', text):
        if word:
            if word[0] not in ",.?!:;)}]»" and prev_word[-1] not in "«[(":
                if first_word:
                    paragraph.add_run('\t')
                else:
                    paragraph.add_run(' ')

            run = paragraph.add_run(word)
            #self.get_color_and_save_word_to_annotation(word, run)
            if formatting == 'b':
                run.bold = True
            if formatting == 'i':
                run.italic = True
            if formatting == 'u':
                run.underline = True
            prev_word = word
            #run.font.size = Pt(self.font_size)
            #run.font.name = self.font_name
    return prev_word

doc = create_colored_docx(doc, "https://ru.wikipedia.org/wiki/%D0%93%D1%80%D0%B0%D1%84_%D0%9D%D0%BE%D1%80%D1%84%D0%BE%D0%BB%D0%BA")
out = io.BytesIO()
doc.save(out)
doc.save('ottt.docx')
doc_bytes = out.getvalue()
print('sent to uno')
conv = client.UnoClient(port=2000)
pdf_bytes = conv.convert(indata=doc_bytes, convert_to='pdf')
images = convert_from_bytes(pdf_bytes)

from PIL import Image
for i, img in enumerate(images):
    img.save(f'img_{i}.png')
